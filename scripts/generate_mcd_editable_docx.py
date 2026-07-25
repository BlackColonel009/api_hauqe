from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT / "scripts"))
from generate_mcd_pdf import PAGES  # noqa: E402


OUTPUT = ROOT / "output" / "docx" / "MCD_HAUQE_CERTIF_EDITABLE.docx"
GREEN = "0B6B4B"
DARK_GREEN = "064C38"
LIGHT_GREEN = "E7F3EE"
GOLD = "D3A229"
LIGHT_GOLD = "FFF4D6"
INK = "1E2925"
MUTED = "64706B"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=110, bottom=80, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_run(run, size=9, bold=False, color=INK, italic=False) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_table_width(table, widths_cm: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            row.cells[idx].width = Cm(width)
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def add_entity_card(cell, entity) -> None:
    cell.text = ""
    set_cell_margins(cell, top=100, start=120, bottom=100, end=120)
    title = cell.paragraphs[0]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run(entity.name)
    set_run(run, size=10, bold=True, color="FFFFFF")
    set_cell_shading(cell, GREEN)

    inner = cell.add_table(rows=1, cols=2)
    inner.style = "Table Grid"
    set_table_width(inner, [5.7, 6.1])
    headers = ("Attribut conceptuel", "Rôle / remarque modifiable")
    for idx, label in enumerate(headers):
        p = inner.rows[0].cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(label), size=8, bold=True, color=DARK_GREEN)
        set_cell_shading(inner.rows[0].cells[idx], LIGHT_GOLD)
        set_cell_margins(inner.rows[0].cells[idx])
    set_repeat_table_header(inner.rows[0])

    for attr in entity.attrs:
        row = inner.add_row()
        prevent_row_split(row)
        is_id = attr.startswith("#")
        p1 = row.cells[0].paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        set_run(p1.add_run(attr), size=8, bold=is_id, color=DARK_GREEN if is_id else INK)
        p2 = row.cells[1].paragraphs[0]
        p2.paragraph_format.space_after = Pt(0)
        set_run(
            p2.add_run("Identifiant conceptuel" if is_id else ""),
            size=8,
            italic=not is_id,
            color=MUTED,
        )
        for inner_cell in row.cells:
            set_cell_margins(inner_cell, top=55, bottom=55)
            inner_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(80)
    p.paragraph_format.space_after = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("MCD HAUQE CERTIF"), size=30, bold=True, color=GREEN)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(26)
    set_run(p.add_run("Modèle conceptuel de données détaillé — version Word modifiable"), size=17, color=DARK_GREEN)

    info = doc.add_table(rows=4, cols=2)
    info.style = "Table Grid"
    set_table_width(info, [5.0, 19.8])
    rows = (
        ("Notation", "Merise / présentation inspirée de PowerDesigner"),
        ("Contenu", "Entités, attributs, identifiants, associations et cardinalités"),
        ("Format", "Tous les tableaux et textes de ce document sont directement modifiables dans Word"),
        ("Statut", "Version 0.2 — brouillon à soumettre à la revue HAUQE/GFA"),
    )
    for row, values in zip(info.rows, rows):
        prevent_row_split(row)
        for idx, value in enumerate(values):
            p = row.cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_run(p.add_run(value), size=10, bold=idx == 0, color=DARK_GREEN if idx == 0 else INK)
            set_cell_shading(row.cells[idx], LIGHT_GREEN if idx == 0 else "FFFFFF")
            set_cell_margins(row.cells[idx], top=120, bottom=120)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(8)
    set_run(p.add_run("Mode d’emploi"), size=15, bold=True, color=GREEN)
    for text in (
        "Modifier les noms et attributs directement dans les cartes d’entités.",
        "Modifier les verbes et cardinalités dans les tableaux d’associations.",
        "Ajouter ou supprimer des lignes selon les arbitrages métier.",
        "Après validation, utiliser cette version pour dériver le MLD et le MPD PostgreSQL.",
    ):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(5)
        set_run(p.add_run(text), size=10)
    doc.add_page_break()


def add_domain(doc: Document, page) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    set_run(p.add_run(f"{page.code} — {page.title}"), size=18, bold=True, color=GREEN)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    set_run(p.add_run(page.subtitle), size=10, italic=True, color=MUTED)

    entities = list(page.entities)
    for start in range(0, len(entities), 2):
        outer = doc.add_table(rows=1, cols=2)
        outer.autofit = False
        outer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_table_width(outer, [12.25, 12.25])
        row = outer.rows[0]
        prevent_row_split(row)
        for idx in range(2):
            if start + idx < len(entities):
                add_entity_card(row.cells[idx], entities[start + idx])
            else:
                row.cells[idx].text = ""
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(3)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(5)
    set_run(p.add_run("Associations et cardinalités"), size=13, bold=True, color=GREEN)

    rel_table = doc.add_table(rows=1, cols=6)
    rel_table.style = "Table Grid"
    set_table_width(rel_table, [4.6, 2.0, 4.3, 4.6, 2.0, 7.0])
    headers = ("Entité source", "Card.", "Association", "Entité cible", "Card.", "Commentaire / arbitrage")
    for idx, label in enumerate(headers):
        cell = rel_table.rows[0].cells[idx]
        set_cell_shading(cell, GREEN)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(label), size=8, bold=True, color="FFFFFF")
    set_repeat_table_header(rel_table.rows[0])

    for rel in page.relations:
        row = rel_table.add_row()
        prevent_row_split(row)
        values = (rel.left, rel.left_card, rel.verb, rel.right, rel.right_card, "")
        for idx, value in enumerate(values):
            cell = row.cells[idx]
            set_cell_margins(cell, top=65, bottom=65)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in (1, 2, 4) else WD_ALIGN_PARAGRAPH.LEFT
            set_run(p.add_run(value), size=7.6, bold=idx == 2, color=DARK_GREEN if idx == 2 else INK)

    doc.add_page_break()


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(9)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for style_name, size, color, before, after in (
        ("Heading 1", 16, GREEN, 18, 10),
        ("Heading 2", 13, GREEN, 14, 7),
        ("Heading 3", 12, DARK_GREEN, 10, 5),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def build() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)
    section.header_distance = Cm(0.55)
    section.footer_distance = Cm(0.55)
    configure_styles(doc)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(header.add_run("HAUQE Certif — MCD éditable"), size=8, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(footer.add_run("Document de travail — Version 0.2 — 23 juillet 2026"), size=8, color=MUTED)

    add_cover(doc)
    for page in PAGES:
        add_domain(doc, page)
    if doc.paragraphs and not doc.paragraphs[-1].text:
        p = doc.paragraphs[-1]
        p._element.getparent().remove(p._element)

    doc.core_properties.title = "MCD HAUQE Certif — version Word modifiable"
    doc.core_properties.subject = "Modèle conceptuel de données détaillé"
    doc.core_properties.author = "Projet HAUQE Certif"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
